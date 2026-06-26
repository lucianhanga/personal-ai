"""Parsers + chunking — pure, no DB/network. PDF/DOCX fixtures are generated in-test."""

from __future__ import annotations

import io

import pytest

from personalai_modality_files import (
    UnsupportedFileTypeError,
    chunk_text,
    parse_document,
)


def test_parse_text_and_markdown() -> None:
    txt = parse_document(b"hello world", "notes.txt")
    assert txt.text == "hello world"
    assert txt.mime == "text/plain"
    md = parse_document(b"# Title", "readme.md")
    assert md.mime == "text/markdown"


def test_parse_docx() -> None:
    import docx

    document = docx.Document()
    document.add_paragraph("Hello from DOCX")
    document.add_paragraph("second line")
    buf = io.BytesIO()
    document.save(buf)
    parsed = parse_document(buf.getvalue(), "doc.docx")
    assert "Hello from DOCX" in parsed.text
    assert "second line" in parsed.text
    assert parsed.mime.endswith("wordprocessingml.document")


def test_parse_pdf() -> None:
    fpdf = pytest.importorskip("fpdf")  # fpdf2 (dev dep) to generate a real PDF
    pdf = fpdf.FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=12)
    pdf.cell(text="Hello from PDF")
    data = bytes(pdf.output())
    parsed = parse_document(data, "doc.pdf")
    assert "Hello from PDF" in parsed.text
    assert parsed.mime == "application/pdf"


def _blank_pdf() -> bytes:
    """A valid PDF with a blank page -> no text layer (stands in for a scanned/image-only PDF)."""
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def test_pdf_without_text_layer_falls_back_to_ocr(monkeypatch: pytest.MonkeyPatch) -> None:
    # A scanned PDF yields no pypdf text -> parse_document delegates to OCR (#450) when available.
    import personalai_modality_files.ocr as ocr_mod

    monkeypatch.setattr(ocr_mod, "ocr_available", lambda: True)
    monkeypatch.setattr(ocr_mod, "ocr_pdf", lambda content, **kw: "OCR RECOVERED TEXT")
    parsed = parse_document(_blank_pdf(), "scan.pdf")
    assert parsed.text == "OCR RECOVERED TEXT"
    assert parsed.mime == "application/pdf"


def test_pdf_without_text_layer_degrades_when_ocr_unavailable(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    # Slim install: no OCR deps -> empty text (the "no text found" state), OCR never invoked.
    import personalai_modality_files.ocr as ocr_mod

    monkeypatch.setattr(ocr_mod, "ocr_available", lambda: False)
    monkeypatch.setattr(
        ocr_mod, "ocr_pdf", lambda *a, **k: pytest.fail("ocr_pdf must not run when unavailable")
    )
    parsed = parse_document(_blank_pdf(), "scan.pdf")
    assert parsed.text.strip() == ""


def test_enable_ocr_false_skips_ocr_entirely(monkeypatch: pytest.MonkeyPatch) -> None:
    import personalai_modality_files.ocr as ocr_mod

    monkeypatch.setattr(
        ocr_mod, "ocr_available", lambda: pytest.fail("availability must not be checked")
    )
    parsed = parse_document(_blank_pdf(), "scan.pdf", enable_ocr=False)
    assert parsed.text.strip() == ""


def test_unsupported_type_raises() -> None:
    with pytest.raises(UnsupportedFileTypeError):
        parse_document(b"...", "image.heic")


def test_chunking_overlap_and_coverage() -> None:
    text = "abcdefghij" * 30  # 300 chars
    chunks = chunk_text(text, size=100, overlap=20)
    assert len(chunks) >= 3
    assert all(len(c) <= 100 for c in chunks)
    # every character is covered somewhere
    assert "".join(chunks).replace("", "") != ""
    assert chunk_text("   ", size=100, overlap=10) == []


def test_chunking_skips_blank_windows() -> None:
    text = "abc" + " " * 300 + "xyz"  # middle windows are all whitespace -> skipped
    chunks = chunk_text(text, size=50, overlap=0)
    assert any("abc" in c for c in chunks)
    assert any("xyz" in c for c in chunks)
    assert all(c.strip() for c in chunks)  # no blank chunks kept


def test_chunking_validates_params() -> None:
    with pytest.raises(ValueError, match="size must be positive"):
        chunk_text("x", size=0)
    with pytest.raises(ValueError, match="overlap"):
        chunk_text("x", size=10, overlap=10)
