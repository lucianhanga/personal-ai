"""Lightweight file parsers + chunking for ingestion (M3-2).

Pure functions over in-memory bytes (no network, no DB). Supports txt/markdown (direct), PDF
(pypdf), and DOCX (python-docx). Treat all input as untrusted: we only extract text — no active
content is executed.
"""

from __future__ import annotations

import io
from dataclasses import dataclass
from pathlib import PurePosixPath

_TEXT_EXTS = {".txt": "text/plain", ".md": "text/markdown", ".markdown": "text/markdown"}


@dataclass(frozen=True)
class ParsedDocument:
    """Extracted text plus the detected MIME type.

    ``ocr`` is True when the text came from the OCR fallback (a scanned / image-only PDF, #450)
    rather than an embedded text layer; ``pages`` is the PDF page count when known. Both let the UI
    surface a truthful "OCR'd N pages" step in the document pipeline activity timeline.
    """

    text: str
    mime: str
    ocr: bool = False
    pages: int | None = None


class UnsupportedFileTypeError(Exception):
    """Raised when a file extension is not supported."""


def parse_document(content: bytes, filename: str, *, enable_ocr: bool = True) -> ParsedDocument:
    """Extract text from ``content`` based on ``filename``'s extension.

    For a PDF with no embedded text layer (a scanned / image-only document) ``pypdf`` yields
    nothing; when ``enable_ocr`` is set and the optional OCR dependencies are installed, fall
    back to OCR (#450). OCR is slow (CPU-bound, ~0.6s/page), so async callers should invoke
    ``parse_document`` via ``asyncio.to_thread`` to avoid blocking the event loop.
    """
    ext = PurePosixPath(filename).suffix.lower()
    if ext in _TEXT_EXTS:
        return ParsedDocument(text=content.decode("utf-8", errors="replace"), mime=_TEXT_EXTS[ext])
    if ext == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        pages = len(reader.pages)
        text = "\n\n".join((page.extract_text() or "") for page in reader.pages)
        ocr_used = False
        if not text.strip() and enable_ocr:
            # No text layer: a scanned / image-only PDF. OCR it when available (#450); otherwise
            # leave the text empty so the caller shows the "no text found" empty state.
            from personalai_modality_files.ocr import ocr_available, ocr_pdf

            if ocr_available():
                text = ocr_pdf(content)
                ocr_used = bool(text.strip())
        return ParsedDocument(text=text, mime="application/pdf", ocr=ocr_used, pages=pages)
    if ext == ".docx":
        import docx

        document = docx.Document(io.BytesIO(content))
        text = "\n".join(p.text for p in document.paragraphs)
        return ParsedDocument(
            text=text,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    raise UnsupportedFileTypeError(f"unsupported file type: {ext or filename!r}")


def chunk_text(text: str, *, size: int = 1000, overlap: int = 150) -> list[str]:
    """Split ``text`` into overlapping chunks of up to ``size`` characters."""
    if size <= 0:
        raise ValueError("size must be positive")
    if not 0 <= overlap < size:
        raise ValueError("overlap must be >= 0 and < size")
    cleaned = text.strip()
    if not cleaned:
        return []
    step = size - overlap
    chunks: list[str] = []
    start = 0
    while start < len(cleaned):
        chunk = cleaned[start : start + size].strip()
        if chunk:
            chunks.append(chunk)
        start += step
    return chunks
